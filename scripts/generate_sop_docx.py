import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_sop_docx(filename="PropertyWorks_Admin_SOP_Guide.docx"):
    doc = docx.Document()
    
    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles Colors
    NAVY = RGBColor(0, 27, 79)      # #001B4F
    GOLD = RGBColor(212, 161, 58)    # #D4A13A
    DARK = RGBColor(30, 41, 59)      # #1E293B
    SLATE = RGBColor(71, 85, 105)    # #475569

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_t1 = p_title.add_run("PropertyWorks ")
    run_t1.font.name = "Arial"
    run_t1.font.size = Pt(24)
    run_t1.font.bold = True
    run_t1.font.color.rgb = NAVY

    run_t2 = p_title.add_run("Website Administration Guide & SOP")
    run_t2.font.name = "Arial"
    run_t2.font.size = Pt(24)
    run_t2.font.bold = True
    run_t2.font.color.rgb = GOLD

    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run("Client Standard Operating Procedure (SOP) & Content Management Handbook")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = SLATE

    # Divider Line
    p_div = doc.add_paragraph()
    p_div_run = p_div.add_run("―" * 55)
    p_div_run.font.color.rgb = GOLD
    p_div_run.font.bold = True

    # Callout Box Helper
    def add_callout(text, title="IMPORTANT NOTICE"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.8)
        
        shading = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(r'''
            <w:tcBorders {} >
                <w:top w:val="none"/>
                <w:left w:val="single" w:sz="24" w:space="0" w:color="001B4F"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
            </w:tcBorders>
        '''.format(nsdecls('w')))
        tcPr.append(borders)
        
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after = Pt(4)
        
        r_t = cp.add_run(f"{title}\n")
        r_t.font.name = "Arial"
        r_t.font.size = Pt(10)
        r_t.font.bold = True
        r_t.font.color.rgb = NAVY
        
        r_b = cp.add_run(text)
        r_b.font.name = "Arial"
        r_b.font.size = Pt(9.5)
        r_b.font.color.rgb = DARK
        doc.add_paragraph()

    add_callout(
        "This handbook provides step-by-step guidance on managing PropertyWorks website content, real estate projects, "
        "articles, contact channels, lead consent, and admin security settings. Everything updated inside the visual Admin Panel "
        "reflects live on the website instantly.",
        "WELCOME & OVERVIEW"
    )

    # Section 1
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = NAVY

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = GOLD

    def add_body(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.name = "Arial"
            r_pre.font.size = Pt(10)
            r_pre.font.bold = True
            r_pre.font.color.rgb = DARK
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.color.rgb = DARK

    # 1. Admin Access & Portal Overview
    add_h1("1. Admin Access & Portal Details")
    add_body("You can access the PropertyWorks Administrator Dashboard from any browser using the following details:")

    t1 = doc.add_table(rows=3, cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Environment", "Admin Panel URL", "Master Password"]
    for i, h in enumerate(headers):
        cell = t1.cell(0, i)
        cell.text = h
        shading = parse_xml(r'<w:shd {} w:fill="001B4F"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(255, 255, 255)

    data1 = [
        ["Live Production Site", "https://www.propertyworks.in/admin", "admin123"],
        ["Local Development", "http://localhost:5173/admin", "admin123"]
    ]
    for row_idx, row_data in enumerate(data1, start=1):
        for col_idx, text in enumerate(row_data):
            cell = t1.cell(row_idx, col_idx)
            cell.text = text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = DARK

    doc.add_paragraph()

    # 2. How to Change Admin Password
    add_h1("2. How to Change the Admin Password (via Code Base)")
    add_body("For security reasons, the administrator login password is stored inside the database configuration file.")
    add_body(" Follow these exact steps to update your Admin password:", "Step-by-Step Security Instructions:")
    
    add_body("Open your project root directory in VS Code or your preferred code editor.", "Step 1: ")
    add_body("Open the configuration file located at: backend/database/admin.json", "Step 2: ")
    add_body("Locate the 'passwordHash' key and change its value to your new password:", "Step 3: ")

    # Code Box Table
    t_code = doc.add_table(rows=1, cols=1)
    t_code.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_code = t_code.cell(0, 0)
    c_code.width = Inches(6.8)
    shading = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
    c_code._tc.get_or_add_tcPr().append(shading)
    cp = c_code.paragraphs[0]
    r_code = cp.add_run('{\n  "passwordHash": "YourNewPasswordHere123"\n}')
    r_code.font.name = "Courier New"
    r_code.font.size = Pt(10)
    r_code.font.bold = True
    r_code.font.color.rgb = NAVY

    doc.add_paragraph()
    add_body("Save the file. Your new password is now active immediately for all future logins!", "Step 4: ")

    # 3. SOP: Website Content & Images (via Admin Panel)
    add_h1("3. SOP: Managing Website Content & Images (via Admin Panel)")
    add_body("You can easily manage all real estate projects, articles, contact channels, lead consent, and page sections directly from the visual Admin Panel.")

    add_h2("A. Real Estate Projects Management")
    add_body("Go to the Admin Panel -> Projects tab. Click the + Add New Project button. Enter Project Title, Developer, Location, Budget, Property Type (Residential/Commercial), Overview, and Amenities. Upload project hero images and floor plans. Click Save Project.", "1. Adding a New Project: ")
    add_body("Click the Edit (Pencil) button next to any listed project. Modify prices, specifications, availability, or images, then click Save Project.", "2. Editing an Existing Project: ")
    add_body("Click the Trash icon next to any obsolete project listing to remove it instantly from the website.", "3. Deleting a Project: ")

    add_h2("B. Knowledge Center & Blog Articles")
    add_body("Go to Articles tab -> Click + Add New Article. Enter Title, Category, Author, Read Time, and Markdown Content body.", "1. Publishing Articles: ")
    add_body("Fill out Meta Title and Meta Description for search engine optimization.", "2. SEO Optimization: ")

    add_h2("C. Contact Channels & Legal Consent Wording")
    add_body("Go to Settings tab -> Edit WhatsApp Number or Contact Phone.", "1. Update Phone / WhatsApp: ")
    add_body("Update Contact Email to support@propertyworks.in. This automatically syncs across the Footer, Contact Page, and all Legal Policy pages (Privacy Policy, Terms & Conditions, Disclaimer, Cookie Policy, Refund Policy).", "2. Support Email Sync: ")
    add_body("Modify the legal consent wording displayed in the Shortlist Modals.", "3. Mandatory Lead Consent: ")

    add_h2("D. Section Headlines & Gold Text Formatting")
    add_body("Go to Page Sections tab to edit text headlines, subheadings, quotes, and descriptions for Hero, Services, About Us, Why Choose Us, and FAQs.", "1. Page Sections Customization: ")
    add_body("To highlight any word in golden color on the website, simply wrap the word inside [gold] tags! Example: 'Stop Evaluating Real Estate [gold]Blindly.[/gold]'", "2. Golden Highlight Rule: ")

    # 4. Lead Management
    add_h1("4. Lead Management & Export")
    add_body("All customer inquiries submitted via 'Get My Free Residential Shortlist' or 'Get My Free Commercial Shortlist' modals are recorded in real time.", "Automatic Lead Storage: ")
    add_body("Go to Leads tab in Admin Panel to view client contact numbers, budget, location preferences, and timestamp.", "Viewing Leads: ")
    add_body("Click Export to CSV / Excel to download all customer details for your sales & CRM follow-up team.", "Exporting Leads: ")

    # 5. Google Business Profile Reference
    add_h1("5. PropertyWorks Google Business Profile Quick Reference")
    add_body("Use these verified business details when updating your Google Business Profile (GBP), social accounts, or directory listings:")

    t2 = doc.add_table(rows=8, cols=2)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    gbp_info = [
        ("Business Name", "PropertyWorks"),
        ("Official Website", "https://www.propertyworks.in"),
        ("Primary Phone", "+91-8433826365"),
        ("Support Email", "support@propertyworks.in"),
        ("Business Category", "Real Estate Advisory"),
        ("Service Areas", "Mumbai, Thane, Navi Mumbai and surrounding Growth corridors"),
        ("Operating Hours", "10:00 AM – 6:00 PM (Monday to Saturday)"),
        ("Registered Address", "B-202, Lodha Splendora, G.B Road, Thane West, 400615")
    ]
    for idx, (label, val) in enumerate(gbp_info):
        c1 = t2.cell(idx, 0)
        c2 = t2.cell(idx, 1)
        c1.text = label
        c2.text = val
        
        shading = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
        c1._tc.get_or_add_tcPr().append(shading)
        
        for p in c1.paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = NAVY
        for p in c2.paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9.5)
                r.font.color.rgb = DARK

    doc.add_paragraph()
    add_callout(
        "For technical support, custom features, or server deployment assistance, please contact your website engineering team.",
        "TECHNICAL SUPPORT"
    )

    doc.save(filename)
    print("DOCX generated successfully:", filename)

if __name__ == "__main__":
    create_sop_docx()
